from flask import Flask, request, jsonify, send_from_directory
from flask_cors import CORS
from docx import Document
from docx.shared import RGBColor
import os
import PyPDF2
import fitz  # PyMuPDF
from pdf2docx import Converter
import random

# code begin
import subprocess
import ast
from time import sleep
import random
# code end

app = Flask(__name__)
CORS(app)

UPLOAD_FOLDER = './folder'
PROCESSED_FOLDER = './processed'
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['PROCESSED_FOLDER'] = PROCESSED_FOLDER

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({'error': '没有文件部分'})

    file = request.files['file']
    flag = int(request.form.get('flag', 0))  # 获取 flag 的值，默认为 0
    if file.filename == '':
        return jsonify({'error': '未选择文件'})

    if file:
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
        file.save(filepath)
        if allowed_file_docx(file.filename):
            detected_filepath = os.path.join(app.config['PROCESSED_FOLDER'], 'detected_' + file.filename)
            result=process_docx_file(filepath, detected_filepath, flag)#一段就直接一个概率，文件的话可以凑一个概率？
            return jsonify({'result': result, 'processed_file': 'detected_' + file.filename})
        elif allowed_file_pdf(file.filename):
            word_path = os.path.join(app.config['PROCESSED_FOLDER'], 'output_demo.docx')
            detected_filepath = os.path.join(app.config['PROCESSED_FOLDER'], 'detected_' + file.filename)
            # process_pdf_file(filepath, detected_filepath)
            pdf_to_word_pdf2docx(filepath, word_path)
            result = add_flag_to_pdf(filepath, detected_filepath, word_path, flag)
            # result = 0.85 # 一段就直接一个概率，文件的话可以凑一个概率？
            return jsonify({'result': result, 'processed_file': 'detected_' + file.filename})

def process_docx_file(input_path, output_path, flag):
    doc = Document(input_path)
    ai_probability=0.0
    para_length=0
    ai_length=0
    for para in doc.paragraphs:
        if len(para.text.split()) > 10:
            # 假设这里是检测逻辑
            #ai_probability = detect_ai_text(para.text)
            para_length += len(para.text.split())
            if flag==1:
                ai_probability=just_roberta_pt(para.text)
            else:
                ai_probability = detect_ai_text(para.text)
            print(ai_probability)
            if ai_probability > 0.8:
                ai_length += len(para.text.split())
                for run in para.runs:
                    run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

    doc.save(output_path)
    print(ai_length)
    print(para_length)
    result=ai_length/para_length
    return result

def allowed_file_docx(filename): # 判断是不是docx文件
    return '.' in filename and filename.rsplit('.', 1)[1].lower() == 'docx'

def allowed_file_pdf(filename): # 判断是不是pdf文件
    return '.' in filename and filename.rsplit('.', 1)[1].lower() == 'pdf'

def detect_ai_text(text):
    # 替换为实际检测逻辑
    # code begin
    if os.path.exists("pre.txt") and os.path.exists("ppl.txt") and os.path.exists("Bscore.txt"):
        os.remove("pre.txt")
        os.remove("ppl.txt")
        os.remove("Bscore.txt")
    subprocess.run(["sbatch", "./run_roberta.sh", text])
    subprocess.run(["sbatch", "./run_ppl.sh", text])
    with open('Bscore.txt', 'w') as f:
        f.write(str(random.uniform(0,1)))
        sleep(10)

    sleep(1)
    while 1:
        if os.path.exists("pre.txt") and os.path.exists("ppl.txt") and os.path.exists("Bscore.txt"):
            break
    sleep(1)
    
    if os.path.exists("result.txt"):
        os.remove("result.txt")
    
    subprocess.run(["sbatch", "./run_three.sh"])
    prediction=0.0

    sleep(1)
    while 1:
        if os.path.exists("result.txt"):
            break
    sleep(1)
    
    with open('result.txt', 'r') as f:
        prediction=float("{:.2f}".format(float(f.read())))
    # code end
    return prediction


def pdf_to_word_pdf2docx(pdf_path, word_path):
    cv = Converter(pdf_path)
    cv.convert(word_path, start=0, end=None)
    cv.close()

def detect(text):
    # code begin
    if os.path.exists("pre.txt") and os.path.exists("ppl.txt") and os.path.exists("Bscore.txt"):
        os.remove("pre.txt")
        os.remove("ppl.txt")
        os.remove("Bscore.txt")
    subprocess.run(["sbatch", "./run_roberta.sh", text])
    subprocess.run(["sbatch", "./run_ppl.sh", text])
    with open('Bscore.txt', 'w') as f:
        # f.write(str(random.uniform(0,1)))
        f.write('0.0015102209213308882')
        sleep(10)

    sleep(1)
    while 1:
        if os.path.exists("pre.txt") and os.path.exists("ppl.txt") and os.path.exists("Bscore.txt"):
            break
    sleep(1)

    if os.path.exists("result.txt"):
        os.remove("result.txt")
    
    subprocess.run(["sbatch", "./run_three.sh"])
    prediction=0.0

    sleep(1)
    while 1:
        if os.path.exists("result.txt"):
            break
    sleep(1)

    with open('result.txt', 'r') as f:
        prediction=float("{:.2f}".format(float(f.read())))
    # code end
    return prediction


def add_flag_to_pdf(pdf_path, output_path, word_path, tag):
    # 打开 PDF 文件
    document = fitz.open(pdf_path)
    doc = Document(word_path)
    i = 0
    ai_text=0
    all_text=0

    # 遍历每一页
    for page_num in range(len(document)):
        page = document[page_num]
        # 获取页面中的所有文本块
        text_instances = page.get_text("dict")["blocks"]

        # 遍历每个文本块
        for block in text_instances:
            i += 1
            if i > 2 and i % 2 == 1:
                continue
            # 从第1段开始，每段分别对应i=1, 2, 4, 6
            if i==1:
                para = doc.paragraphs[1].text
            else:
                para = doc.paragraphs[round(i/2)+1].text
            print(para)
            # 此处para为pdf的每一段的内容
            all_text += len(para.split())
            #percentage = detect(para)
            if tag==1:
                percentage=just_roberta_pt(para)
            else:
                percentage = detect(para)
            if(percentage < 0.8):
                continue # 表示该段不是llm生成，不用添加flag
            if(len(para.split()) < 10):
                all_text -= len(para.split())
                continue
            ai_text += len(para.split())
            # 如果执行下面语句表示该段是llm生成，添加flag

            if block['type'] == 0:  # 仅处理文本块
                # 如果文本块中有内容
                if len(block["lines"]) > 0:
                    # 获取第一个文本块中的第一个文本行的第一个文本跨度
                    first_span = block["lines"][0]["spans"][0]
                    # 获取文本块的边界框
                    bbox = fitz.Rect(first_span["bbox"])
                    # 计算插入位置，即段落开头前方一点位置
                    flag_position = bbox.tl
                    flag_position = fitz.Point(flag_position.x - 20, flag_position.y)
                    # 在段落开头插入红色文本 "flag"
                    page.insert_text(flag_position, "llm generated", fontsize=first_span["size"], color=(1, 0, 0))

    # 保存修改后的 PDF 文件，并进行垃圾回收和压缩
    document.save(output_path, garbage=4, deflate=True)
    result=ai_text/all_text
    print(ai_text)
    print(all_text)
    print(result)
    return result


@app.route('/processed/<filename>', methods=['GET'])
def download_file(filename):
    return send_from_directory(app.config['PROCESSED_FOLDER'], filename)

@app.route('/check', methods=['POST'])
def check_content():
    data = request.get_json()
    text = data.get('text', '')
    flag = data.get('flag', 0)  # 获取 flag 的值，默认为 0
    if not text:
        return jsonify({'error': '没有提供文本'})
    if flag == 1:
        #probability = flag_process_text(text)  # 调用快速检测函数
        probability=just_roberta_pt(text)
    else:
        probability = process_text(text)
        # probability=roberta_pt(text)
    return jsonify({'result': probability})

def process_text(text):
    # code begin
    if os.path.exists("pre.txt") and os.path.exists("ppl.txt") and os.path.exists("Bscore.txt"):
        os.remove("pre.txt")
        os.remove("ppl.txt")
        os.remove("Bscore.txt")
    print("get it")
    subprocess.run(["sbatch", "./run_roberta.sh", text])
    subprocess.run(["sbatch", "./run_ppl.sh", text])
    with open('Bscore.txt', 'w') as f:
        f.write(str(random.uniform(0,1)))
        sleep(10)

    sleep(1)
    while 1:
        if os.path.exists("pre.txt") and os.path.exists("ppl.txt") and os.path.exists("Bscore.txt"):
            break
    sleep(1)

    if os.path.exists("result.txt"):
        os.remove("result.txt")
    
    subprocess.run(["sbatch", "./run_three.sh"])
    prediction=0.0

    sleep(1)
    while 1:
        if os.path.exists("result.txt"):
            break
    sleep(1)

    with open('result.txt', 'r') as f:
        prediction=float("{:.2f}".format(float(f.read())))
    # code end
    return prediction


# code begin
def roberta_pt(text):
    if os.path.exists("pre.txt"):
        os.remove("pre.txt")

    subprocess.run(["sbatch", "./run_roberta.sh", text])

    sleep(1)
    while not os.path.exists("pre.txt"):
        sleep(1)
    sleep(5)

    # 打印文件大小
    print(f"File size of 'pre.txt': {os.path.getsize('pre.txt')} bytes")

    roberta_score = 0.0
    with open('pre.txt', 'r') as f:
        content = f.read().strip()
        print(f"Content read from 'pre.txt':\n{content}\n")
        if content:
            try:
                roberta_score = round(ast.literal_eval(content)[0][0], 2)
                print(f"Evaluated score: {roberta_score}")
            except (SyntaxError, ValueError) as e:
                print(f"Error reading 'pre.txt': {e}")
    return roberta_score


def just_roberta_pt(text):
    if os.path.exists("pre.txt"):
        os.remove("pre.txt")

    subprocess.run(["sbatch", "./run_roberta.sh", text])

    sleep(1)
    while not (os.path.exists("pre.txt") and os.path.getsize("pre.txt") > 10):
        sleep(1)

    # 打印文件大小
    print(f"File size of 'pre.txt': {os.path.getsize('pre.txt')} bytes")

    roberta_score = 0.0
    with open('pre.txt', 'r') as f:
        content = f.read().strip()
        print(f"Content read from 'pre.txt':\n{content}\n")
        if content:
            try:
                roberta_score = round(ast.literal_eval(content)[0][0], 2)
                print(f"Evaluated score: {roberta_score}")
            except (SyntaxError, ValueError) as e:
                print(f"Error reading 'pre.txt': {e}")
    return roberta_score



if __name__ == '__main__':
    os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
    os.makedirs(app.config['PROCESSED_FOLDER'], exist_ok=True)
    app.run(host='0.0.0.0', port=5000)
