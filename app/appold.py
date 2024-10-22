from flask import Flask, request, jsonify, send_from_directory
from flask_cors import CORS
from docx import Document
from docx.shared import RGBColor
import os
import PyPDF2
import fitz  # PyMuPDF
from pdf2docx import Converter
import random

app = Flask(__name__)
CORS(app)

UPLOAD_FOLDER = './folder'
PROCESSED_FOLDER = './processed'
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['PROCESSED_FOLDER'] = PROCESSED_FOLDER

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({'error': '没有文件部分'})

    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': '未选择文件'})

    if file:
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
        file.save(filepath)
        if allowed_file_docx(file.filename):
            detected_filepath = os.path.join(app.config['PROCESSED_FOLDER'], 'detected_' + file.filename)
            process_docx_file(filepath, detected_filepath)
            result = 0.85 #一段就直接一个概率，文件的话可以凑一个概率？
            return jsonify({'result': result, 'processed_file': 'detected_' + file.filename})
        elif allowed_file_pdf(file.filename):
            word_path = os.path.join(app.config['PROCESSED_FOLDER'], 'output_demo.docx')
            detected_filepath = os.path.join(app.config['PROCESSED_FOLDER'], 'detected_' + file.filename)
            # process_pdf_file(filepath, detected_filepath)
            pdf_to_word_pdf2docx(filepath, word_path)
            add_flag_to_pdf(filepath, detected_filepath, word_path)
            result = 0.85 # 一段就直接一个概率，文件的话可以凑一个概率？
            return jsonify({'result': result, 'processed_file': 'detected_' + file.filename})

def process_docx_file(input_path, output_path):
    doc = Document(input_path)

    for para in doc.paragraphs:
        if len(para.text.split()) > 10:
            # 假设这里是检测逻辑
            ai_probability = detect_ai_text(para.text)
            if ai_probability > 0.8:
                for run in para.runs:
                    run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

    doc.save(output_path)

def allowed_file_docx(filename): # 判断是不是docx文件
    return '.' in filename and filename.rsplit('.', 1)[1].lower() == 'docx'

def allowed_file_pdf(filename): # 判断是不是pdf文件
    return '.' in filename and filename.rsplit('.', 1)[1].lower() == 'pdf'

def detect_ai_text(text):
    # 替换为实际检测逻辑
    return 0.9

def pdf_to_word_pdf2docx(pdf_path, word_path):
    cv = Converter(pdf_path)
    cv.convert(word_path, start=0, end=None)
    cv.close()

def detect(text):
    return random.choice([0.85, 0.15])
def add_flag_to_pdf(pdf_path, output_path, word_path):
    # 打开 PDF 文件
    document = fitz.open(pdf_path)
    doc = Document(word_path)
    i = 0

    # 遍历每一页
    for page_num in range(len(document)):
        page = document[page_num]
        # 获取页面中的所有文本块
        text_instances = page.get_text("dict")["blocks"]

        # 遍历每个文本块
        for block in text_instances:
            i += 1
            if i > 2 and i % 2 == 1:
                continue
            # 从第1段开始，每段分别对应i=1, 2, 4, 6
            if i==1:
                para = doc.paragraphs[1].text
            else:
                para = doc.paragraphs[round(i/2)+1].text
            # 此处para为pdf的每一段的内容
            percentage = detect(para)
            if(percentage < 0.2):
                continue # 表示该段不是llm生成，不用添加flag
            if(len(para.split()) < 10):
                continue
            # 如果执行下面语句表示该段是llm生成，添加flag

            if block['type'] == 0:  # 仅处理文本块
                # 如果文本块中有内容
                if len(block["lines"]) > 0:
                    # 获取第一个文本块中的第一个文本行的第一个文本跨度
                    first_span = block["lines"][0]["spans"][0]
                    # 获取文本块的边界框
                    bbox = fitz.Rect(first_span["bbox"])
                    # 计算插入位置，即段落开头前方一点位置
                    flag_position = bbox.tl
                    flag_position = fitz.Point(flag_position.x - 20, flag_position.y)
                    # 在段落开头插入红色文本 "flag"
                    page.insert_text(flag_position, "llm generated", fontsize=first_span["size"], color=(1, 0, 0))

    # 保存修改后的 PDF 文件，并进行垃圾回收和压缩
    document.save(output_path, garbage=4, deflate=True)


@app.route('/processed/<filename>', methods=['GET'])
def download_file(filename):
    return send_from_directory(app.config['PROCESSED_FOLDER'], filename)

@app.route('/check', methods=['POST'])
def check_content():
    data = request.get_json()
    text = data.get('text', '')
    if not text:
        return jsonify({'error': '没有提供文本'})
    print("get it")
    probability = process_text(text)
    print(probability)
    return jsonify({'result': probability})

def process_text(text):
    probability = 0.755
    return probability

if __name__ == '__main__':
    os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
    os.makedirs(app.config['PROCESSED_FOLDER'], exist_ok=True)
    app.run(host='0.0.0.0', port=6006)
